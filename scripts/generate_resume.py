#!/usr/bin/env python3
"""Generate a black-and-white, one-page-oriented offerpilot_skill resume DOCX."""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:
    raise SystemExit(
        "缺少 python-docx。请使用 Codex workspace dependencies 提供的 Python 运行此脚本。"
    ) from exc

LATIN_FONT = "Times New Roman"
EAST_ASIA_FONT = "Microsoft YaHei"
BLACK = RGBColor(0, 0, 0)

DENSITY = {
    "compact": {"body": 9.2, "small": 8.8, "name": 17.5, "section": 11.0, "line": 1.0, "gap": 0.5, "margin": 1.15},
    "normal": {"body": 10.0, "small": 9.3, "name": 18.5, "section": 11.5, "line": 1.08, "gap": 1.2, "margin": 1.35},
    "spacious": {"body": 10.8, "small": 10.0, "name": 20.0, "section": 12.0, "line": 1.22, "gap": 2.2, "margin": 1.55},
}


def set_run_font(run, size, bold=False, italic=False):
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key, value in (
        ("ascii", LATIN_FONT),
        ("hAnsi", LATIN_FONT),
        ("cs", LATIN_FONT),
        ("eastAsia", EAST_ASIA_FONT),
    ):
        rfonts.set(qn(f"w:{key}"), value)


def set_paragraph(p, cfg, before=0, after=None, line=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(cfg["gap"] if after is None else after)
    p.paragraph_format.line_spacing = cfg["line"] if line is None else line
    p.paragraph_format.widow_control = True


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Cm(width_cm)


def set_table_width_pct(table, width_pct=100):
    """Use Word's AutoFit-to-Window behavior instead of a fixed table width."""
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_pct * 50)))
    tbl_w.set(qn("w:type"), "pct")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")
    table.autofit = True


def set_cell_width_pct(cell, width_pct):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_pct * 50)))
    tc_w.set(qn("w:type"), "pct")


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, value="nil", color="000000", size="0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), value)
        tag.set(qn("w:color"), color)
        tag.set(qn("w:sz"), size)


def section_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")  # eighths of a point = 1 pt
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def add_text(p, text, cfg, *, size=None, bold=False, italic=False):
    run = p.add_run(str(text))
    set_run_font(run, cfg["body"] if size is None else size, bold=bold, italic=italic)
    return run


def load_manifest(path: Path) -> dict:
    with path.open(encoding="utf-8-sig") as handle:
        data = json.load(handle)
    if not isinstance(data.get("resume"), dict):
        raise ValueError("manifest 缺少 resume object")
    return data


def run_gate(manifest: Path):
    checker = Path(__file__).with_name("quality_check.py")
    env = os.environ.copy()
    env["PYTHONUTF8"] = "1"
    completed = subprocess.run(
        [sys.executable, str(checker), str(manifest), "--json"],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        env=env,
    )
    if completed.returncode:
        raise ValueError("质量闸门未通过，禁止生成终稿：\n" + completed.stdout.strip())


def content_score(resume: dict) -> tuple[int, int]:
    chars = len(str(resume.get("objective", ""))) + len(str(resume.get("availability", "")))
    blocks = 0
    for section in resume.get("sections") or []:
        chars += len(str(section.get("title", "")))
        for entry in section.get("entries") or []:
            blocks += 1
            chars += sum(len(str(entry.get(k, ""))) for k in ("title", "meta", "subtitle", "background"))
            for bullet in entry.get("bullets") or []:
                if isinstance(bullet, dict):
                    chars += len(str(bullet.get("title", ""))) + len(str(bullet.get("text", "")))
                else:
                    chars += len(str(bullet))
        for item in section.get("items") or []:
            blocks += 1
            if isinstance(item, dict):
                chars += len(str(item.get("category", item.get("title", "")))) + len(str(item.get("content", "")))
            else:
                chars += len(str(item))
    return chars, blocks


def choose_density(resume: dict, requested: str) -> str:
    if requested != "auto":
        return requested
    chars, blocks = content_score(resume)
    if chars > 1900 or blocks > 9:
        return "compact"
    if chars < 950 or blocks < 6:
        return "spacious"
    return "normal"


def configure_document(doc, cfg):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(cfg["margin"])
        section.left_margin = section.right_margin = Cm(1.45)
    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(cfg["body"])
        style.font.color.rgb = BLACK
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), LATIN_FONT)
        rfonts.set(qn("w:hAnsi"), LATIN_FONT)
        rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def add_placeholder(cell, cfg):
    nested = cell.add_table(rows=1, cols=1)
    set_table_width_pct(nested)
    set_table_borders(nested, value="single", color="000000", size="8")
    pcell = nested.cell(0, 0)
    set_cell_width_pct(pcell, 100)
    pcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row = nested.rows[0]
    row.height = Cm(2.6)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    p = pcell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, cfg, after=0)
    add_text(p, "证件照", cfg, size=cfg["small"])


def add_header(doc, resume, cfg, manifest_dir: Path):
    # Equal left/right photo-width columns keep the information block centered
    # on the physical page while the photo stays anchored at the top-right.
    table = doc.add_table(rows=1, cols=3)
    set_table_width_pct(table)
    set_table_borders(table)
    spacer, info, right = table.rows[0].cells
    set_cell_width_pct(spacer, 13)
    set_cell_width_pct(info, 74)
    set_cell_width_pct(right, 13)
    set_cell_margins(spacer, end=0)
    set_cell_margins(info, start=40, end=40)
    set_cell_margins(right, start=80)
    info.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, cfg, after=2)
    add_text(p, resume["name"], cfg, size=cfg["name"], bold=True)

    p = info.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, cfg, after=2)
    add_text(p, resume["contact"], cfg, size=cfg["small"])

    objective = str(resume.get("objective", "")).strip()
    availability = str(resume.get("availability", "")).strip()
    if objective or availability:
        p = info.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, cfg, after=0)
        if objective:
            add_text(p, objective, cfg, size=cfg["body"], bold=True)
        if objective and availability:
            add_text(p, "  |  ", cfg, size=cfg["body"])
        if availability:
            add_text(p, availability, cfg, size=cfg["body"])

    photo = str(resume.get("photo_path", "")).strip()
    if photo:
        photo_path = Path(photo).expanduser()
        if not photo_path.is_absolute():
            photo_path = manifest_dir / photo_path
        if not photo_path.exists():
            raise ValueError(f"证件照文件不存在：{photo_path}")
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph(p, cfg, after=0)
        p.add_run().add_picture(str(photo_path), height=Cm(2.6))
    else:
        right.paragraphs[0]._element.getparent().remove(right.paragraphs[0]._element)
        add_placeholder(right, cfg)


def add_section_title(doc, title, cfg):
    p = doc.add_paragraph()
    set_paragraph(p, cfg, before=4 if cfg is DENSITY["spacious"] else 3, after=1)
    add_text(p, title, cfg, size=cfg["section"], bold=True)
    section_rule(p)


def add_entry_heading(doc, entry, cfg):
    p = doc.add_paragraph()
    set_paragraph(p, cfg, before=2, after=0.5)
    if entry.get("meta"):
        p.paragraph_format.tab_stops.add_tab_stop(Cm(17.7), WD_TAB_ALIGNMENT.RIGHT)
    add_text(p, entry["title"], cfg, bold=True)
    if entry.get("subtitle"):
        add_text(p, " | " + str(entry["subtitle"]), cfg, bold=True)
    if entry.get("meta"):
        add_text(p, "\t" + str(entry["meta"]), cfg, size=cfg["small"])


def add_background(doc, background, cfg):
    p = doc.add_paragraph()
    set_paragraph(p, cfg, after=0.5)
    add_text(p, "项目背景  ", cfg, bold=True)
    add_text(p, background, cfg)


def add_bullet(doc, bullet, cfg, require_title=True):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, cfg, after=0.5)
    p.paragraph_format.left_indent = Cm(0.48)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    if isinstance(bullet, dict):
        title = str(bullet.get("title", "")).strip()
        if title:
            add_text(p, title + "  ", cfg, bold=True)
        elif require_title:
            raise ValueError("工作/项目 bullet 缺少 title")
        add_text(p, bullet.get("text", ""), cfg)
    else:
        if require_title:
            raise ValueError("工作/项目 bullet 必须为含 title/text 的 object")
        add_text(p, bullet, cfg)


def add_items(doc, items, cfg):
    for item in items:
        if not isinstance(item, dict):
            raise ValueError("技能/自我评价 item 必须为 object")
        title = str(item.get("category", item.get("title", ""))).strip()
        content = str(item.get("content", "")).strip()
        if not title or not content:
            raise ValueError("技能/自我评价 item 必须包含小标题和内容")
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, cfg, after=0.5)
        p.paragraph_format.left_indent = Cm(0.48)
        p.paragraph_format.first_line_indent = Cm(-0.28)
        add_text(p, title + "  ", cfg, bold=True)
        add_text(p, content, cfg)


def build(data: dict, output: Path, manifest_path: Path, density: str):
    resume = data["resume"]
    for key in ("name", "contact", "sections"):
        if not resume.get(key):
            raise ValueError(f"resume.{key} 缺失，禁止用占位符生成终稿")
    chosen = choose_density(resume, density)
    cfg = DENSITY[chosen]
    doc = Document()
    configure_document(doc, cfg)
    add_header(doc, resume, cfg, manifest_path.parent)

    for section in resume["sections"]:
        section_type = section.get("type", "")
        entries = section.get("entries") or []
        items = section.get("items") or []
        if section_type == "experience" and not entries:
            continue
        if not entries and not items:
            continue
        add_section_title(doc, section["title"], cfg)
        if section_type in {"skills", "self_evaluation"}:
            add_items(doc, items, cfg)
            continue
        require_background = section_type in {"experience", "projects"}
        for entry in entries:
            add_entry_heading(doc, entry, cfg)
            if require_background:
                background = str(entry.get("background", "")).strip()
                if not background:
                    raise ValueError(f"{entry.get('title', '经历')} 缺少项目背景")
                add_background(doc, background, cfg)
            for bullet in entry.get("bullets") or []:
                add_bullet(doc, bullet, cfg, require_title=require_background)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return chosen


def main() -> int:
    parser = argparse.ArgumentParser(description="生成 offerpilot_skill 黑白一页式 DOCX 简历")
    parser.add_argument("manifest", type=Path)
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        help="输出路径；默认保存到 Skill 安装目录下的 output",
    )
    parser.add_argument("--density", choices=("auto", "compact", "normal", "spacious"), default="auto")
    parser.add_argument("--skip-gate", action="store_true", help="仅供调试；正式交付不得使用")
    args = parser.parse_args()
    try:
        data = load_manifest(args.manifest)
        if not args.skip_gate:
            run_gate(args.manifest)
        resume = data.get("resume") or {}
        if args.output:
            output = args.output
        else:
            output_dir = Path(__file__).resolve().parent.parent / "output"
            name = str(resume.get("name") or "候选人").strip()
            role = str(resume.get("objective") or "目标岗位").strip()
            safe = re.sub(r'[<>:"/\\|?*]+', "-", f"{name}-{role}-定向简历").strip(". ")
            output = output_dir / f"{safe}.docx"
        chosen = build(data, output, args.manifest.resolve(), args.density)
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(f"[FAIL] {exc}", file=sys.stderr)
        return 1
    print(f"Generated: {output}")
    print(f"Density: {chosen}")
    print("Next: render with the documents skill; verify one page, photo height, fonts, rules, overflow and whitespace.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
